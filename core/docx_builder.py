"""将生成的论文内容组装为格式规范的 Word 文档"""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

def set_chinese_font(run, font_name="宋体", size=Pt(12)):
    run.font.name = "Times New Roman"
    run.font.size = size
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font_name)

def add_cover(doc, title, author="", school="", date=""):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_chinese_font(run, "黑体", Pt(22))
    run.bold = True

    if author:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"作者：{author}")
        set_chinese_font(run2, "宋体", Pt(14))

    if school:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(school)
        set_chinese_font(run3, "宋体", Pt(14))

    doc.add_page_break()

def markdown_to_docx(doc, md_text, charts=None):
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("### "):
            heading = doc.add_heading(line[4:], level=3)
            for run in heading.runs:
                set_chinese_font(run, "黑体", Pt(14))
        elif line.startswith("## "):
            heading = doc.add_heading(line[3:], level=2)
            for run in heading.runs:
                set_chinese_font(run, "黑体", Pt(15))
        elif line.startswith("# "):
            heading = doc.add_heading(line[2:], level=1)
            for run in heading.runs:
                set_chinese_font(run, "黑体", Pt(16))
                run.bold = True
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            for run in p.runs:
                set_chinese_font(run, "宋体", Pt(12))
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(text, style="List Number")
            for run in p.runs:
                set_chinese_font(run, "宋体", Pt(12))
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5

            bold_parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in bold_parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    set_chinese_font(run, "宋体", Pt(12))
                else:
                    run = p.add_run(part)
                    set_chinese_font(run, "宋体", Pt(12))
        i += 1

    if charts:
        for chart_path in charts:
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def build_thesis(sections: dict, metadata: dict, charts: list = None) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.line_spacing = 1.5

    title = metadata.get("title", "毕业论文")
    author = metadata.get("author", "")
    school = metadata.get("school", "")

    add_cover(doc, title, author, school)

    if "abstract" in sections:
        markdown_to_docx(doc, sections["abstract"])
        doc.add_page_break()

    chapter_order = ["literature", "methodology", "results", "conclusion"]
    chapter_names = {
        "literature": "文献综述",
        "methodology": "研究方法",
        "results": "数据分析与结果",
        "conclusion": "结论与讨论",
    }

    for key in chapter_order:
        if key in sections:
            if not sections[key].startswith("#"):
                heading = doc.add_heading(chapter_names.get(key, key), level=1)
                for run in heading.runs:
                    set_chinese_font(run, "黑体", Pt(16))
                    run.bold = True
            markdown_to_docx(doc, sections[key], charts if key == "results" else None)
            doc.add_page_break()

    output_path = os.path.join("output", f"{title[:20]}.docx")
    doc.save(output_path)
    return output_path
